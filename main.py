from dotenv import load_dotenv
import streamlit as st
from streamlit_mic_recorder import mic_recorder
import openai
from openai import OpenAI
from gtts import gTTS
from pydub import AudioSegment
import base64
from io import BytesIO
import tempfile
import os
import json
import datetime
from langchain_openai import ChatOpenAI
from langchain.schema import HumanMessage, AIMessage, SystemMessage
import uuid
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import bcrypt
from flask import Flask
from flask_wtf import FlaskForm
from wtforms import StringField, PasswordField, SubmitField
from wtforms.validators import DataRequired, Email
from firebase_admin.exceptions import FirebaseError
import firebase_admin
from firebase_admin import credentials, auth, firestore, exceptions

from lib.domain.User import User
from lib.domain.InterviewContext import InterviewContext

# テンプレートの定義は変更なし
# テンプレートの定義
templates = {
    "personal_attributes": """
        あなたは{theme}についてインタビューを行うインタビュアーです。
        回答者の基本的なプロフィールを収集します。
        これまでの会話コンテキスト: {context}
        質問生成の指針:
        - まず，デモグラフィック情報（年齢、職業、家族構成は必須）を聞く
        - 音楽の好みやリスニング習慣に関する質問を含める
        - {theme}の利用経験について尋ねる
        - １対話で複数の質問を投げかけない
        - １対話につき質問は１つとする
        - このフェーズでは合計5つの質問を行う
        - 相槌は最低限で
    """,
    "usage_situation": """
        あなたは{theme}についてインタビューを行うインタビュアーです。
        {theme}の利用状況や消費シーンについて詳しく探ります。
        これまでの会話コンテキスト: {context}
        質問生成の指針:
        - {theme}をどのような場面で利用するか，具体的なエピソードなどを交えて
        - 利用した時の満足点と不満点について，具体的なエピソードなどを交えて
        - {theme}を利用する際の感情や期待を，具体的なエピソードなどを交えて
        - {theme}を利用するにあたりこんな機能があれば，みたいな要望を，具体的なエピソードなどを交えて
        - 各対話につき質問は１つに絞る
        - なぜそう思ったのかを深掘りする（それは本当に他社で満たせていないことですか？）
        - このフェーズでは合計5つの質問を行う
        - 相槌は最低限で
    """,
    "purchase_intention": """
        あなたは{theme}についてインタビューを行うインタビュアーです。
        {theme}の選択意思決定プロセスについて深掘りします。
        これまでの会話コンテキスト: {context}
        質問生成の指針:
        - 選択時に重視する要素（価格、楽曲数、音質など）を聞き，なぜそれを重視するのか深掘りする
        - 選択のきっかけや情報源を具体的に聞く
        - 選択後の満足度や不満を具体的に聞く
        - 各対話につき質問は１つに絞る
        - なぜそう思ったのかを深掘りする（それは本当に他社で満たせていないことですか？）
        - このフェーズでは合計5つの質問を行う
        - 相槌は最低限で
    """,
    "competitor_analysis": """
        あなたは{theme}についてインタビューを行うインタビュアーです。
        競合サービスに対する認識を調査します。
        これまでの会話コンテキスト: {context}
        質問生成の指針:
        - 知っている競合サービスやその特徴を，具体的なエピソードなどを交えて
        - 競合サービスとの比較や選択理由を，具体的なエピソードなどを交えて
        - 競合サービスに対する印象や期待を，具体的なエピソードなどを交えて
        - 各対話につき質問は１つに絞る
        - なぜそう思ったのかを深掘りする（それは本当に他社で満たせていないことですか？）
        - このフェーズでは合計5つの質問を行う
        - 相槌は最低限で
    """,
    "summary": """
        テーマ: {theme}
        インタビュー全体を分析し、以下の形式で分析レポートを作成してください：
        1. どんな{theme}が選ばれるか:
        2. 今の{theme}を選んだ理由:
        3. 他社{theme}と比較したときの魅力:
        4. これから{theme}を選ぶとしたらどこを重視するか:
        5. {theme}における不満や問題:
        6. {theme}において新しく求める特徴や機能:
        これまでの会話コンテキスト: {context}
    """
}

class LoginForm(FlaskForm):
    email = StringField('Email', validators=[DataRequired(), Email()])
    password = PasswordField('Password', validators=[DataRequired()])
    submit = SubmitField('Log In')

def initialize_firestore(credential_path):
    if not firebase_admin._apps:
        cred = credentials.Certificate(credential_path)
        firebase_admin.initialize_app(cred)
    return firestore.client()

def setup_openai_api():
    if not firebase_admin._apps:
        cred = credentials.ApplicationDefault()
        firebase_admin.initialize_app(cred, {
            'projectId': os.environ.get('GOOGLE_CLOUD_PROJECT'),
        })
    # ローカル環境用の設定ファイルを読み込む
    if os.path.exists('.runtimeconfig.json'):
        with open('.runtimeconfig.json', 'r') as config_file:
            config = json.load(config_file)
        api_key = config.get('openai', {}).get('apikey')
    else:
        # 環境変数からAPIキーを取得
        api_key = os.environ.get('OPENAI_API_KEY')
    if not api_key:
        raise ValueError("OpenAI API key is not set.")
    # OpenAIライブラリを使用する場合
    openai.api_key = api_key

def initialize_session_state():
    if "logged_in" not in st.session_state:
        st.session_state.logged_in = False
    if "messages" not in st.session_state:
        st.session_state.messages = []
        st.session_state.current_question = "こんにちは！音楽サブスクサービスについてのインタビューを始めましょう。"
        st.session_state.ai_response_audio_html = ""
        st.session_state.last_question_displayed = False
        st.session_state.phase = "personal_attributes"
        st.session_state.context = ""
        st.session_state.question_count = 0
        st.session_state.theme = "音楽サブスクサービス"

def speech_to_text(audio_bytes):
    client = OpenAI()
    with tempfile.NamedTemporaryFile(delete=False, suffix=".mp3") as temp_audio:
        temp_audio.write(audio_bytes)
        temp_audio_path = temp_audio.name
    with open(temp_audio_path, "rb") as audio_file:
        response = client.audio.transcriptions.create(
            model="whisper-1",
            file=audio_file
        )
    transcript = response.text
    return transcript

def get_ai_response(messages):
    llm = ChatOpenAI(model_name="gpt-4", temperature=0.7)
    response = llm.invoke(messages)
    return response.content

def text_to_speech(text):
    tts = gTTS(text=text, lang='ja')
    audio_io = BytesIO()
    tts.write_to_fp(audio_io)
    audio_io.seek(0)
    audio = AudioSegment.from_file(audio_io, format="mp3")
    faster_audio = audio.speedup(playback_speed=1.2)
    buffer = BytesIO()
    faster_audio.export(buffer, format="mp3")
    audio_base64 = base64.b64encode(buffer.getvalue()).decode()
    return f'<audio autoplay="true" src="data:audio/mp3;base64,{audio_base64}"></audio>'

def save_interview_results(theme, summary, messages):
    try:
        doc = docx.Document()
        title = doc.add_heading(f'{theme} インタビュー結果', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_style = doc.styles['Title']
        title_style.font.size = Pt(24)
        for i, item in enumerate(summary.split('\n')):
            if item:
                paragraph = doc.add_paragraph(f"{i+1}. {item}")
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        now = datetime.datetime.now()
        timestamp = now.strftime("%Y%m%d_%H%M%S")
        filename = f"{theme}_interview_results_{timestamp}.docx"
        doc.save(filename)
        uid = firestoreUser.uid

        # メッセージを辞書のリストに変換
        contexts = [
            f"{'ユーザー' if isinstance(msg, HumanMessage) else 'AI'}: {msg.content}"
            for msg in messages
        ]
        interviewId = str(uuid.uuid4())
        # InterviewContextインスタンスの作成
        interview_context = InterviewContext(
            contexts=contexts,
            interviewId=interviewId,
            theme=theme,
            timestamp=timestamp
        )
        # Firestoreに保存
        firestore_client.collection("users").document(uid).collection("interview_contexts").document(interviewId).set(interview_context.to_dict())

        # Firestoreにユーザー情報を更新
        user_ref = firestore_client.collection("users").document(uid)
        user_ref.update({
            "survey_id_list": firestore.ArrayUnion([interviewId])
        })

        

        return filename
    except Exception as e:
        st.error(f"インタビュー結果の保存中にエラーが発生しました: {e}")
        return None

def update_phase():
    if st.session_state.question_count <= 5:
        st.session_state.phase = "personal_attributes"
    elif st.session_state.question_count <= 10:
        st.session_state.phase = "usage_situation"
    elif st.session_state.question_count <= 15:
        st.session_state.phase = "purchase_intention"
    elif st.session_state.question_count <= 20:
        st.session_state.phase = "competitor_analysis"
    else:
        st.session_state.phase = "summary"

def process_user_input(audio):
    user_input = speech_to_text(audio["bytes"])
    st.session_state.messages.append(HumanMessage(content=user_input))
    st.session_state.context += f"ユーザー: {user_input}\n"
    st.session_state.question_count += 1

def generate_ai_response():
    if st.session_state.phase != "summary":
        system_message = SystemMessage(content=templates[st.session_state.phase].format(theme=st.session_state.theme, context=st.session_state.context))
        messages = [system_message] + st.session_state.messages
        ai_response = get_ai_response(messages)
        st.session_state.messages.append(AIMessage(content=ai_response))
        st.session_state.context += f"AI: {ai_response}\n"
        st.session_state.current_question = ai_response
    else:
        summary = get_ai_response([SystemMessage(content=templates["summary"].format(theme=st.session_state.theme, context=st.session_state.context))])
        st.markdown("インタビューを終了します。以下がJOBインサイトの分析結果です：")
        st.markdown(summary)
        filename = save_interview_results(st.session_state.theme, summary, st.session_state.messages)
        if filename:
            st.success(f"インタビュー結果を{filename}に保存しました。")
        st.markdown("インタビューが完了しました。ありがとうございました。")
        ai_response = "インタビューが完了しました。ありがとうございました。"
    return ai_response

def display_form():
    st.title("音楽サブスクサービス インタビュー対話アプリ")
    age = st.number_input("年齢", min_value=0, max_value=120)
    gender = st.selectbox("性別", ["選択してください", "男性", "女性", "その他"])
    occupation = st.text_input("ご職業")
    all_fields_filled = age > 0 and gender != "選択してください" and occupation

    with st.form("interview_form"):
        submit_button = st.form_submit_button("インタビューを開始", disabled=not all_fields_filled)
    
    if submit_button and all_fields_filled:
        st.session_state.form_submitted = True
        st.session_state.context += f"ユーザー情報: 年齢={age}, 性別={gender}, 職業={occupation}\n"
        st.session_state.interview_started = True

        # Firestoreにユーザー情報を更新
        user_ref = firestore_client.collection("users").document(firestoreUser.uid)

        # 年齢と性別をFirestoreに格納
        user_ref.update({
            "age": age,
            "gender": gender
        })

        # セッション状態のユーザー情報も更新
        st.session_state.user.age = age
        st.session_state.user.gender = gender

    if not all_fields_filled:
        st.warning("全ての項目を入力してください。")

# インタビューを実行
def conduct_interview():
    st.title("音楽サブスクサービス インタビュー対話アプリ")

    # 最後の質問が表示されていない場合、現在の質問を表示
    if not st.session_state.last_question_displayed:
        st.markdown(f"<h2 style='text-align: center;'>{st.session_state.current_question}</h2>", unsafe_allow_html=True)

    # 音声入力の受付
    audio = mic_recorder(start_prompt="録音開始", stop_prompt="録音停止", key='recorder', just_once=True)

    # 音声入力がある場合の処理
    if audio and audio.get("bytes"):
        with st.spinner("処理中1/3..."):
            process_user_input(audio)

        update_phase()

        with st.spinner("処理中2/3..."):
            ai_response = generate_ai_response()

        with st.spinner("処理中3/3..."):
            st.session_state.ai_response_audio_html = text_to_speech(ai_response)

        # 最後の質問表示フラグを設定し、ページを再読み込み
        st.session_state.last_question_displayed = True
        st.rerun()

    # 最後の質問が表示された後の処理
    if st.session_state.last_question_displayed:
        st.markdown(f"<h2 style='text-align: center;'>{st.session_state.current_question}</h2>", unsafe_allow_html=True)
        st.markdown(st.session_state.ai_response_audio_html, unsafe_allow_html=True)
        st.session_state.last_question_displayed = False


def firebase_initialization(cred_path):
    if not firebase_admin._apps:
        cred = credentials.Certificate(cred_path)
        firebase_admin.initialize_app(cred, {
            'databaseURL': 'https://n1-interview-default-rtdb.firebaseio.com/',
            'projectId': 'n1-interview',
        })
    return firestore.client()

def display_login_form():
    st.title("ログイン")
    email = st.text_input("メールアドレス")
    password = st.text_input("パスワード", type="password")
    if st.button("ログイン"):
        if email and password:
            success, message = login(email, password)
            if success:
                st.session_state.logged_in = True
                st.session_state.user_id = firestoreUser.uid
                st.success(message)
                st.rerun()
            else:
                st.error(message)
        else:
            st.error("メールアドレスとパスワードを入力してください。")
    st.write("アカウントをお持ちでない方は、こちらからサインアップしてください。")
    if st.button("サインアップ"):
        st.session_state.page = "signup"
        st.rerun()

def initialize_firebase():
    global firestore_client
    if not firebase_admin._apps:
        cred_path = os.environ.get('GOOGLE_APPLICATION_CREDENTIALS')
        cred = credentials.Certificate(cred_path)
        firebase_admin.initialize_app(cred, {
            'databaseURL': 'https://n1-interview-default-rtdb.firebaseio.com/',
            'projectId': 'n1-interview',
        })
    firestore_client = firestore.client()

def getFirestoreUser(uid):
    global firestoreUser
    current_user_doc = firestore_client.collection("users").document(uid).get()
    user_data = current_user_doc.to_dict()
    firestoreUser = User.from_dict(user_data)

def login(email, password):
    try:
        user = auth.get_user_by_email(email)
        # パスワードの検証（実際のアプリケーションではセキュアな方法で行う必要があります）
        custom_token = auth.create_custom_token(user.uid)

        return True, "ログインに成功しました。"
    except auth.UserNotFoundError:
        return False, "メールアドレスまたはパスワードが間違っています。", None
    except Exception as e:
        return False, f"エラーが発生しました: {str(e)}", None


def display_signup_form():
    st.title("サインアップ")
    username = st.text_input("お名前")
    email = st.text_input("メールアドレス")
    password = st.text_input("パスワード", type="password")
    confirm_password = st.text_input("パスワード（確認）", type="password")
    
    if st.button("アカウント作成"):
        if not username or not email or not password or not confirm_password:
            st.error("すべての項目を入力してください。")
        elif password != confirm_password:
            st.error("パスワードが一致しません。")
        else:
            try:
                # Firebase Authenticationでユーザーを作成
                auth_user = auth.create_user(email=email, password=password)
                # Userクラスのインスタンスを作成
                user = User(
                    name=username,
                    age="",
                    uid=auth_user.uid,
                    occupation="",
                    email=email,
                    survey_id_list=[],
                )
                # Firestoreにユーザー情報を追加
                firestore_client.collection("users").document(user.uid).set(user.to_dict())
                
                st.success("アカウントが正常に作成されました。")
                st.session_state.page = "login"
                st.rerun()
            except Exception as e:
                st.error(f"アカウント作成中にエラーが発生しました: {str(e)}")
    st.write("すでにアカウントをお持ちの方は、こちらからログインしてください。")
    if st.button("ログインページへ"):
        st.session_state.page = "login"
        st.rerun()

def main():
    initialize_firebase()  # Firebaseの初期化
    setup_openai_api()
    initialize_session_state()

    if "page" not in st.session_state:
        st.session_state.page = "login"
    if "logged_in" not in st.session_state:
        st.session_state.logged_in = False

    if st.session_state.logged_in:
        if 'form_submitted' not in st.session_state:
            st.session_state.form_submitted = False
        if not st.session_state.form_submitted:
            display_form()
        elif st.session_state.interview_started:
            conduct_interview()
        else:
            st.title("音楽サブスクサービス インタビュー対話アプリ")
            st.write("インタビューの準備ができました。開始ボタンを押してください。")
            if st.button("インタビューを開始"):
                st.session_state.interview_started = True
                st.rerun()
    elif st.session_state.page == "login":
        display_login_form()
    elif st.session_state.page == "signup":
        display_signup_form()

if __name__ == "__main__":
    main()
