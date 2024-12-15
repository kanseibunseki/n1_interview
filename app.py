from dotenv import load_dotenv
import streamlit as st
from streamlit_mic_recorder import mic_recorder
import openai
from gtts import gTTS
from pydub import AudioSegment
import base64
from io import BytesIO
import tempfile
import os
from langchain.chat_models import ChatOpenAI
from langchain.schema import HumanMessage, AIMessage, SystemMessage
import datetime
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

# OpenAI APIキーの設定
# 環境変数からAPIキーを取得
load_dotenv()
api_key = os.getenv("OPENAI_API_KEY")
if not api_key:
    raise ValueError("OPENAI_API_KEY is not set. Please check your .env file.")
openai.api_key = api_key


# セッション状態の初期化
if "messages" not in st.session_state:
    st.session_state.messages = []
    st.session_state.current_question = "こんにちは！音楽サブスクサービスについてのインタビューを始めましょう。"
    st.session_state.ai_response_audio_html = ""
    st.session_state.last_question_displayed = False
    st.session_state.phase = "personal_attributes"
    st.session_state.context = ""
    st.session_state.question_count = 0
    st.session_state.theme = "音楽サブスクサービス"

# テンプレートの定義
templates = {
    "personal_attributes": """
    あなたは{theme}についてインタビューを行うインタビュアーです。
    回答者の基本的なプロフィールを収集します。
    これまでの会話コンテキスト: {context}
    質問生成の指針:
    - まず，デモグラフィック情報（年齢、職業、家族構成は必須）を聞く
    - 音楽の好みやリスニング習慣に関する質問を含める
    - {theme}の利用経験について尋ねる
    - １対話で複数の質問を投げかけない
    - １対話につき質問は１つとする
    - このフェーズでは合計5つの質問を行う
    - 相槌は最低限で
    """,
    "usage_situation": """
    あなたは{theme}についてインタビューを行うインタビュアーです。
    {theme}の利用状況や消費シーンについて詳しく探ります。
    これまでの会話コンテキスト: {context}
    質問生成の指針:
    - {theme}をどのような場面で利用するか，具体的なエピソードなどを交えて
    - 利用した時の満足点と不満点について，具体的なエピソードなどを交えて
    - {theme}を利用する際の感情や期待を，具体的なエピソードなどを交えて
    - {theme}を利用するにあたりこんな機能があれば，みたいな要望を，具体的なエピソードなどを交えて
    - 各対話につき質問は１つに絞る
    - なぜそう思ったのかを深掘りする（それは本当に他社で満たせていないことですか？）
    - このフェーズでは合計5つの質問を行う
    - 相槌は最低限で
    """,
    "purchase_intention": """
    あなたは{theme}についてインタビューを行うインタビュアーです。
    {theme}の選択意思決定プロセスについて深掘りします。
    これまでの会話コンテキスト: {context}
    質問生成の指針:
    - 選択時に重視する要素（価格、楽曲数、音質など）を聞き，なぜそれを重視するのか深掘りする
    - 選択のきっかけや情報源を具体的に聞く
    - 選択後の満足度や不満を具体的に聞く
    - 各対話につき質問は１つに絞る
    - なぜそう思ったのかを深掘りする（それは本当に他社で満たせていないことですか？）
    - このフェーズでは合計5つの質問を行う
    - 相槌は最低限で
    """,
    "competitor_analysis": """
    あなたは{theme}についてインタビューを行うインタビュアーです。
    競合サービスに対する認識を調査します。
    これまでの会話コンテキスト: {context}
    質問生成の指針:
    - 知っている競合サービスやその特徴を，具体的なエピソードなどを交えて
    - 競合サービスとの比較や選択理由を，具体的なエピソードなどを交えて
    - 競合サービスに対する印象や期待を，具体的なエピソードなどを交えて
    - 各対話につき質問は１つに絞る
    - なぜそう思ったのかを深掘りする（それは本当に他社で満たせていないことですか？）
    - このフェーズでは合計5つの質問を行う
    - 相槌は最低限で
    """,
    "summary": """
    テーマ: {theme}
    インタビュー全体を分析し、以下の形式で分析レポートを作成してください：
    1. どんな{theme}が選ばれるか:
    2. 今の{theme}を選んだ理由:
    3. 他社{theme}と比較したときの魅力:
    4. これから{theme}を選ぶとしたらどこを重視するか:
    5. {theme}における不満や問題:
    6. {theme}において新しく求める特徴や機能:
    これまでの会話コンテキスト: {context}
    """
}

def speech_to_text(audio_bytes):
    with tempfile.NamedTemporaryFile(delete=False, suffix=".mp3") as temp_audio:
        temp_audio.write(audio_bytes)
        temp_audio_path = temp_audio.name
    with open(temp_audio_path, "rb") as audio_file:
        transcript = openai.Audio.transcribe("whisper-1", audio_file)
    return transcript["text"]

def get_ai_response(messages):
    llm = ChatOpenAI(model_name="gpt-4", temperature=0.7)
    response = llm(messages)
    return response.content

def text_to_speech(text):
    tts = gTTS(text=text, lang='ja')
    audio_io = BytesIO()
    tts.write_to_fp(audio_io)
    audio_io.seek(0)
    audio = AudioSegment.from_file(audio_io, format="mp3")
    faster_audio = audio.speedup(playback_speed=1.2)
    buffer = BytesIO()
    faster_audio.export(buffer, format="mp3")
    audio_base64 = base64.b64encode(buffer.getvalue()).decode()
    return f'<audio autoplay="true" src="data:audio/mp3;base64,{audio_base64}"></audio>'

def save_interview_results(theme, summary, context):
    try:
        doc = docx.Document()
        title = doc.add_heading(f'{theme} インタビュー結果', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_style = doc.styles['Title']
        title_style.font.size = Pt(24)
        for i, item in enumerate(summary.split('\n')):
            if item:
                paragraph = doc.add_paragraph(f"{i+1}. {item}")
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{theme}_interview_results_{timestamp}.docx"
        doc.save(filename)
        return filename
    except Exception as e:
        st.error(f"インタビュー結果の保存中にエラーが発生しました: {e}")
        return None

st.title("音楽サブスクサービス インタビュー対話アプリ")

if not st.session_state.last_question_displayed:
    st.markdown(f"<h2 style='text-align: center;'>{st.session_state.current_question}</h2>", unsafe_allow_html=True)

audio = mic_recorder(start_prompt="録音開始", stop_prompt="録音停止", key='recorder', just_once=True)

if audio and audio.get("bytes"):
    with st.spinner("Processing your response..."):
        user_input = speech_to_text(audio["bytes"])
        st.session_state.messages.append(HumanMessage(content=user_input))
        st.session_state.context += f"ユーザー: {user_input}\n"
        st.session_state.question_count += 1

    if st.session_state.question_count <= 5:
        st.session_state.phase = "personal_attributes"
    elif st.session_state.question_count <= 10:
        st.session_state.phase = "usage_situation"
    elif st.session_state.question_count <= 15:
        st.session_state.phase = "purchase_intention"
    elif st.session_state.question_count <= 20:
        st.session_state.phase = "competitor_analysis"
    else:
        st.session_state.phase = "summary"

    with st.spinner("Thinking..."):
        if st.session_state.phase != "summary":
            system_message = SystemMessage(content=templates[st.session_state.phase].format(theme=st.session_state.theme, context=st.session_state.context))
            messages = [system_message] + st.session_state.messages
            ai_response = get_ai_response(messages)
            st.session_state.messages.append(AIMessage(content=ai_response))
            st.session_state.context += f"AI: {ai_response}\n"
            st.session_state.current_question = ai_response
        else:
            summary = get_ai_response([SystemMessage(content=templates["summary"].format(theme=st.session_state.theme, context=st.session_state.context))])
            st.markdown("インタビューを終了します。以下がJOBインサイトの分析結果です：")
            st.markdown(summary)
            filename = save_interview_results(st.session_state.theme, summary, st.session_state.context)
            if filename:
                st.success(f"インタビュー結果を{filename}に保存しました。")
            st.markdown("インタビューが完了しました。ありがとうございました。")
            ai_response = "インタビューが完了しました。ありがとうございました。"

    with st.spinner("Generating audio response..."):
        st.session_state.ai_response_audio_html = text_to_speech(ai_response)

    st.session_state.last_question_displayed = True
    st.rerun()

if st.session_state.last_question_displayed:
    st.markdown(f"<h2 style='text-align: center;'>{st.session_state.current_question}</h2>", unsafe_allow_html=True)
    st.markdown(st.session_state.ai_response_audio_html, unsafe_allow_html=True)
    st.session_state.last_question_displayed = False
