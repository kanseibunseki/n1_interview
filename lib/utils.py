import os
import tempfile
from io import BytesIO
import base64
import datetime
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import openai
from gtts import gTTS
from pydub import AudioSegment

# OpenAI APIキーの設定
openai.api_key = os.environ.get("OPENAI_API_KEY")

def speech_to_text(audio_bytes):
    with tempfile.NamedTemporaryFile(delete=False, suffix=".mp3") as temp_audio:
        temp_audio.write(audio_bytes)
        temp_audio_path = temp_audio.name

    with open(temp_audio_path, "rb") as audio_file:
        transcript = openai.Audio.transcribe("whisper-1", audio_file)
    os.remove(temp_audio_path) # 不要になった一時ファイルを削除
    return transcript["text"]

def text_to_speech(text):
    tts = gTTS(text=text, lang='ja')
    audio_io = BytesIO()
    tts.write_to_fp(audio_io)
    audio_io.seek(0)

    audio = AudioSegment.from_file(audio_io, format="mp3")
    faster_audio = audio.speedup(playback_speed=1.2)

    buffer = BytesIO()
    faster_audio.export(buffer, format="mp3")
    audio_base64 = base64.b64encode(buffer.getvalue()).decode()

    return f'<audio autoplay="true" src="data:audio/mp3;base64,{audio_base64}"></audio>'

async def save_interview_results(theme, summary, context):
    try:
        doc = docx.Document()
        title = doc.add_heading(f'{theme} インタビュー結果', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_style = doc.styles['Title']
        title_style.font.size = Pt(24)
        for i, item in enumerate(summary.split('\n')):
            if item:
                paragraph = doc.add_paragraph(f"{i+1}. {item}")
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{theme}_interview_results_{timestamp}.docx"
        doc.save(filename)
        return filename
    except Exception as e:
        print(f"インタビュー結果の保存中にエラーが発生しました: {e}")
        return None

def get_ai_response(messages):
    response = openai.ChatCompletion.create(
        model="gpt-4",
        messages=messages
    )
    return response.choices[0].message["content"]