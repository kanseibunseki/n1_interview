import asyncio
import os
import csv
import datetime
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from langchain.llms import OpenAI
from langchain.chat_models import ChatOpenAI
import chainlit as cl
from langchain.schema import AIMessage, HumanMessage, SystemMessage
import traceback
import openai
import requests
import json
import sounddevice as sd
import numpy as np
import wave
import io

@cl.password_auth_callback
def auth_callback(username: str, password: str):
    if (username, password) == ("admin", "admin"):
        return cl.User(
            identifier="admin",
            metadata={"role": "admin", "provider": "credentials"}
        )
    else:
        return None

templates = {
    "personal_attributes": """
    あなたは{theme}についてインタビューを行うインタビュアーです。
    回答者の基本的なプロフィールを収集します。

    これまでの会話コンテキスト: {context}

    質問生成の指針:
    - まず，デモグラフィック情報（年齢、職業、家族構成は必須）を聞く
    - ライフスタイルや価値観に関する質問を含める
    - {theme}に関連する趣味や習慣について尋ねる
    - １対話で複数の質問を投げかけない
    - １対話につき質問は１つとする
    """,

    "usage_situation": """
    あなたは{theme}についてインタビューを行うインタビュアーです。
    {theme}の利用状況や消費シーンについて詳しく探ります。

    これまでの会話コンテキスト: {context}

    質問生成の指針:
    - {theme}をどのような場面で利用するか，具体的なエピソードなどを交えて
    - 利用した時の満足と不満について，具体的なエピソードなどを交えて
    - {theme}を利用する際の感情や期待を，具体的なエピソードなどを交えて
    - {theme}を利用するにあたりこんなものがあれば，みたいな要望を，具体的なエピソードなどを交えて
    - 各対話につき質問は１つに絞る
    - なぜそう思ったのかを深掘りする
    - 必要があれば，「それは他のお店でも満たしていないニーズでしょうか？」と確認する
    """,

    "purchase_intention": """
    あなたは{theme}についてインタビューを行うインタビュアーです。
    {theme}の選択意思決定プロセスについて深掘りします。

    これまでの会話コンテキスト: {context}

    質問生成の指針:
    - 選択時に重視する要素（価格、品質、ブランドなど）を聞き，なぜそれを重視するのか深掘りする
    - 選択のきっかけや情報源を具体的に聞く
    - 選択後の満足度や不満を具体的に聞く
    - 各対話につき質問は１つに絞る
    - なぜそう思ったのかを深掘りする
    - 必要があれば，「それは他のお店でも満たしていないニーズでしょうか？」と確認する
    """,

    "competitor_analysis": """
    あなたは{theme}についてインタビューを行うインタビュアーです。
    競合製品やブランドに対する認識を調査します。

    これまでの会話コンテキスト: {context}

    質問生成の指針:
    - 知っている競合ブランドやその特徴を，具体的なエピソードなどを交えて
    - 競合サービスとの比較や選択理由を，具体的なエピソードなどを交えて
    - 競合サービスに対する印象や期待を，具体的なエピソードなどを交えて
    - 各対話につき質問は１つに絞る
    - なぜそう思ったのかを深掘りする
    """,

    "summary": """
    テーマ: {theme}

    インタビュー全体を分析し、以下の形式で分析レポートを作成してください：

    1. どんな{theme}が選ばれるか:
    2. 今の{theme}を選んだ理由:
    3. 他社{theme}と比較したときの魅力:
    4. これから{theme}を選ぶとしたらどこを重視するか:
    5. {theme}における不満や問題:
    6. {theme}において新しく求める特徴や機能:

    これまでの会話コンテキスト: {context}
    """
}

async def save_interview_results(theme, summary, context):
    try:
        doc = docx.Document()
        title = doc.add_heading(f'{theme} インタビュー結果', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_style = doc.styles['Title']
        title_style.font.size = Pt(24)
        for i, item in enumerate(summary.split('\n')):
            if item:
                paragraph = doc.add_paragraph(f"{i+1}. {item}")
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{theme}_interview_results_{timestamp}.docx"
        doc.save(filename)
        return filename
    except Exception as e:
        print(f"インタビュー結果の保存中にエラーが発生しました: {e}")
        return None

async def speech_to_text():
    try:
        # 音声録音
        duration = 5  # 録音時間（秒）
        sample_rate = 44100  # サンプリングレート
        channels = 1  # モノラル

        print("録音を開始します...")
        recording = sd.rec(int(duration * sample_rate), samplerate=sample_rate, channels=channels)
        sd.wait()
        print("録音が完了しました。")

        # 録音データをWAVファイルに変換
        wav_buffer = io.BytesIO()
        with wave.open(wav_buffer, 'wb') as wav_file:
            wav_file.setnchannels(channels)
            wav_file.setsampwidth(2)  # 16-bit
            wav_file.setframerate(sample_rate)
            wav_file.writeframes((recording * 32767).astype(np.int16).tobytes())

        # OpenAI Whisper APIを使用して音声をテキストに変換
        wav_buffer.seek(0)
        transcript = openai.Audio.transcribe("whisper-1", wav_buffer, language="ja")
        
        return transcript['text']
    except Exception as e:
        print(f"音声認識中にエラーが発生しました: {e}")
        return None

async def text_to_speech(text):
    try:
        # VOICEVOXのエンドポイント
        url = "http://localhost:50021"

        # 音声合成のリクエスト
        params = {
            "text": text,
            "speaker": 1
        }
        response = requests.post(f"{url}/audio_query", params=params)
        response.raise_for_status()
        query = response.json()

        # 音声合成の実行
        response = requests.post(f"{url}/synthesis", params=params, json=query)
        response.raise_for_status()

        # 音声ファイルの保存
        with open("response.wav", "wb") as f:
            f.write(response.content)

        # 音声ファイルの再生（この部分は環境によって異なる場合があります）
        os.system("aplay response.wav")

    except Exception as e:
        print(f"音声合成中にエラーが発生しました: {e}")

@cl.on_chat_start
async def start():
    try:
        theme = "音楽サブスクサービス"
        purpose = f"リ音楽サブスクサービスを選んだ理由とリ音楽サブスクサービスへの不満やニーズをもとに，どのような{theme}が選ばれるのかを調査したい"
        attributes = f"{theme}をすでに利用したことがあるユーザ"

        session_data = {
            "messages": [],
            "phase": "personal_attributes",
            "context": "",
            "question_count": 0,
            "theme": theme,
            "purpose": purpose,
            "attributes": attributes
        }
        cl.user_session.set("session_data", session_data)

        llm = ChatOpenAI(model_name="gpt-4", temperature=0.7)

        intro_message = "本インタビューは商品開発に消費者様の生の声を生かすことを目的にしています。お答えいただいた内容は匿名化されるのでご安心ください。"
        await cl.Message(content=intro_message).send()
        await text_to_speech(intro_message)

        first_question = await llm.apredict(
            templates["personal_attributes"].format(theme=theme, context="まずなんとお呼びしたら良いですか？のみ質問してください，全体を通して１対話に複数質問をすることを禁じます"))
        await cl.Message(content=first_question).send()
        await text_to_speech(first_question)
    except Exception as e:
        error_message = f"エラーが発生しました: {str(e)}\n"
        error_message += traceback.format_exc()
        await cl.Message(content=error_message).send()

@cl.on_message
async def main(message: cl.Message):
    try:
        session_data = cl.user_session.get("session_data")
        if not session_data:
            raise ValueError("セッションデータが見つかりません")

        messages = session_data["messages"]
        phase = session_data["phase"]
        context = session_data["context"]
        question_count = session_data["question_count"]
        theme = session_data["theme"]

        # 音声入力を使用
        user_input = await speech_to_text()
        if user_input is None:
            await cl.Message(content="音声を認識できませんでした。もう一度お試しください。").send()
            await text_to_speech("音声を認識できませんでした。もう一度お試しください。")
            return

        user_message = HumanMessage(content=user_input)
        messages.append(user_message)
        context += f"ユーザー: {user_input}\n"

        llm = ChatOpenAI(model_name="gpt-4", temperature=0.7)

        question_count += 1

        if question_count == 1:
            phase = "personal_attributes"
            await cl.Message(content="現在のフェーズ: プロフィール").send()
            await text_to_speech("現在のフェーズ: プロフィール")
        elif question_count == 4:
            phase = "usage_situation"
            await cl.Message(content="現在のフェーズ: 利用状況の把握").send()
            await text_to_speech("現在のフェーズ: 利用状況の把握")
        elif question_count == 9:
            phase = "purchase_intention"
            await cl.Message(content="現在のフェーズ: 購入意思の把握").send()
            await text_to_speech("現在のフェーズ: 購入意思の把握")
        elif question_count == 16:
            phase = "competitor_analysis"
            await cl.Message(content="現在のフェーズ: 競合調査").send()
            await text_to_speech("現在のフェーズ: 競合調査")
        else:  # question_count >= 5
            summary = await llm.apredict(templates["summary"].format(theme=theme, context=context))
            await cl.Message(content="インタビューを終了します。以下がJOBインサイトの分析結果です：").send()
            await text_to_speech("インタビューを終了します。以下がJOBインサイトの分析結果です：")
            await cl.Message(content=summary).send()
            await text_to_speech(summary)

            filename = await save_interview_results(theme, summary, context)
            await cl.Message(content=f"インタビュー結果を{filename}に保存しました。").send()
            await text_to_speech(f"インタビュー結果を{filename}に保存しました。")
            await cl.Message(content="インタビューが完了しました。ありがとうございました。").send()
            await text_to_speech("インタビューが完了しました。ありがとうございました。")
            return

        response = await llm.apredict(templates[phase].format(theme=theme, context=context))
        bot_message = AIMessage(content=response)
        messages.append(bot_message)
        context += f" {response}\n"

        await cl.Message(content=response).send()
        await text_to_speech(response)

        session_data.update({
            "messages": messages,
            "phase": phase,
            "context": context,
            "question_count": question_count
        })
        cl.user_session.set("session_data", session_data)

    except Exception as e:
        error_message = f"エラーが発生しました: {str(e)}\n"
        error_message += traceback.format_exc()
        await cl.Message(content=error_message).send()
        await text_to
