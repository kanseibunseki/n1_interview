import datetime

import streamlit as st
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

class DocumentSaver:
    """Handles saving interview results to a document"""
    @staticmethod
    def save_interview_results(theme, summary, context):
        try:
            doc = docx.Document()
            
            # Title
            title = doc.add_heading(f'{theme} インタビュー結果', level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_style = doc.styles['Title']
            title_style.font.size = Pt(24)
            
            # Summary points
            for i, item in enumerate(summary.split('\n'), 1):
                if item:
                    paragraph = doc.add_paragraph(f"{i}. {item}")
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Save file
            timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"{theme}_interview_results_{timestamp}.docx"
            doc.save(filename)
            return filename
        except Exception as e:
            st.error(f"インタビュー結果の保存中にエラーが発生しました: {e}")
            return None
